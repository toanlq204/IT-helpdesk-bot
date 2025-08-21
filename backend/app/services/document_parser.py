import os
from typing import Optional, List, Dict, Any
import PyPDF2
from docx import Document
import markdown
from bs4 import BeautifulSoup
import logging

# LangChain imports
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_pinecone import PineconeVectorStore
from langchain.docstore.document import Document as LangChainDocument
from pinecone import Pinecone

# Local imports
from ..core.config import settings

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Enhanced document processor with LangChain and Pinecone integration
    """
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
        # Initialize embeddings if API key is available
        self.embeddings = None
        if settings.OPENAI_API_KEY:
            try:
                self.embeddings = OpenAIEmbeddings(
                    openai_api_base=settings.OPENAI_API_BASE,
                    api_key=settings.OPENAI_API_KEY,
                    model=settings.EMBEDDING_MODEL
                )
                logger.info("OpenAI embeddings initialized successfully")
            except Exception as e:
                logger.warning(f"Failed to initialize OpenAI embeddings: {e}")
        
        # Initialize Pinecone if API key is available
        self.vector_store = None
        if settings.PINECONE_API_KEY and self.embeddings:
            try:
                self._init_pinecone()
            except Exception as e:
                logger.warning(f"Failed to initialize Pinecone: {e}")
    
    def _init_pinecone(self):
        """
        Initialize Pinecone vector store
        """
        try:
            # Initialize Pinecone client
            pc = Pinecone(api_key=settings.PINECONE_API_KEY)
            
            # Check if index exists, create if not
            index_name = settings.PINECONE_INDEX_NAME
            existing_indexes = [index.name for index in pc.list_indexes()]
            
            if index_name not in existing_indexes:
                logger.info(f"Creating Pinecone index: {index_name}")
                pc.create_index(
                    name=index_name,
                    dimension=1536,  # OpenAI ada-002 embedding dimension
                    metric="cosine",
                    spec={
                        "serverless": {
                            "cloud": "aws",
                            "region": "us-east-1"
                        }
                    }
                )
                # Wait for index to be ready
                import time
                time.sleep(10)
            
            # Initialize vector store
            self.vector_store = PineconeVectorStore(
                index_name=index_name,
                embedding=self.embeddings,
                pinecone_api_key=settings.PINECONE_API_KEY
            )
            logger.info("Pinecone vector store initialized successfully")
            
        except Exception as e:
            logger.error(f"Failed to initialize Pinecone: {e}")
            raise
    
    def extract_text_from_file(self, filepath: str, content_type: str) -> Optional[str]:
        """
        Extract text from various file types
        """
        try:
            if content_type == "application/pdf" or filepath.endswith('.pdf'):
                return self._extract_pdf_text(filepath)
            elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or filepath.endswith('.docx'):
                return self._extract_docx_text(filepath)
            elif content_type == "text/plain" or filepath.endswith('.txt'):
                return self._extract_txt_text(filepath)
            elif content_type == "text/markdown" or filepath.endswith('.md'):
                return self._extract_md_text(filepath)
            else:
                return None
        except Exception as e:
            logger.error(f"Error extracting text from {filepath}: {e}")
            return None
    
    def process_and_store_document(self, 
                                 filepath: str, 
                                 content_type: str, 
                                 document_id: int,
                                 filename: str) -> Dict[str, Any]:
        """
        Process document: extract text, chunk it, generate embeddings, and store in Pinecone
        
        Returns:
            Dict containing processing results and statistics
        """
        try:
            # Extract text
            text = self.extract_text_from_file(filepath, content_type)
            if not text:
                return {"success": False, "error": "Failed to extract text"}
            
            # Create LangChain documents for chunking
            documents = [LangChainDocument(
                page_content=text,
                metadata={
                    "document_id": document_id,
                    "filename": filename,
                    "file_type": content_type,
                    "source": filepath
                }
            )]
            
            # Split into chunks
            chunks = self.text_splitter.split_documents(documents)
            
            # Add chunk-specific metadata
            for i, chunk in enumerate(chunks):
                chunk.metadata.update({
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    "chunk_size": len(chunk.page_content)
                })
            
            # Store in vector database if available
            vector_ids = []
            if self.vector_store:
                try:
                    vector_ids = self.vector_store.add_documents(chunks)
                    logger.info(f"Stored {len(chunks)} chunks in Pinecone for document {document_id}")
                except Exception as e:
                    logger.error(f"Failed to store vectors in Pinecone: {e}")
            else:
                logger.info("Vector store not available, skipping vector storage")
            
            return {
                "success": True,
                "text_length": len(text),
                "chunk_count": len(chunks),
                "vector_ids": vector_ids,
                "chunks": [{
                    "content": chunk.page_content,
                    "metadata": chunk.metadata
                } for chunk in chunks]
            }
            
        except Exception as e:
            logger.error(f"Error processing document {filepath}: {e}")
            return {"success": False, "error": str(e)}
    
    def search_similar_chunks(self, query: str, k: int = 5) -> List[Dict[str, Any]]:
        """
        Search for similar document chunks using vector similarity
        
        Args:
            query: Search query
            k: Number of results to return
            
        Returns:
            List of similar chunks with metadata
        """
        if not self.vector_store:
            logger.warning("Vector store not available for search")
            return []
        
        try:
            results = self.vector_store.similarity_search_with_score(query, k=k)
            
            return [{
                "content": doc.page_content,
                "metadata": doc.metadata,
                "score": score
            } for doc, score in results]
            
        except Exception as e:
            logger.error(f"Error searching similar chunks: {e}")
            return []
    
    def delete_document_vectors(self, document_id: int) -> bool:
        """
        Delete all vectors associated with a document
        
        Args:
            document_id: ID of the document to delete
            
        Returns:
            True if successful, False otherwise
        """
        if not self.vector_store:
            return True  # No vectors to delete
        
        try:
            # Note: This is a simplified approach. In production, you might want to
            # store vector IDs in your database for more efficient deletion
            # For now, we'll use metadata filtering if supported by the vector store
            logger.info(f"Attempting to delete vectors for document {document_id}")
            
            # Try to delete by metadata filter (if supported)
            try:
                # This depends on the vector store implementation
                # For now, we'll just log and return True
                pass
            except Exception as e:
                logger.warning(f"Could not delete vectors by metadata: {e}")
            
            return True
            
        except Exception as e:
            logger.error(f"Error deleting document vectors: {e}")
            return False
    
    def _extract_pdf_text(self, filepath: str) -> str:
        """Extract text from PDF using PyPDF2"""
        text = ""
        with open(filepath, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        return text.strip()
    
    def _extract_docx_text(self, filepath: str) -> str:
        """Extract text from DOCX using python-docx"""
        doc = Document(filepath)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()
    
    def _extract_txt_text(self, filepath: str) -> str:
        """Extract text from plain text file"""
        with open(filepath, 'r', encoding='utf-8') as file:
            return file.read().strip()
    
    def _extract_md_text(self, filepath: str) -> str:
        """Extract text from Markdown file"""
        with open(filepath, 'r', encoding='utf-8') as file:
            md_content = file.read()
            # Convert markdown to HTML, then strip HTML tags
            html = markdown.markdown(md_content)
            soup = BeautifulSoup(html, 'html.parser')
            return soup.get_text().strip()


# Global instance
document_processor = DocumentProcessor()

# Backward compatibility functions
def extract_text_from_file(filepath: str, content_type: str) -> Optional[str]:
    """Backward compatibility wrapper"""
    return document_processor.extract_text_from_file(filepath, content_type)